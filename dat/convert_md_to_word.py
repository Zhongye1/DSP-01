import os
import re
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor


def extract_questions_from_md(file_path):
    """
    从markdown文件中提取题目
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 分离选择题和应用题
    parts = content.split('## 二、综合应用题')
    choice_questions_text = parts[0]
    application_questions_text = parts[1] if len(parts) > 1 else ""
    
    # 提取选择题
    choice_pattern = r'(\d+)\.\s*(.*?)(?=\n\d+\.|\n##|$)'
    choice_matches = re.findall(choice_pattern, choice_questions_text, re.DOTALL)
    
    choice_questions = []
    for match in choice_matches:
        question_num = match[0]
        question_content = match[1].strip()
        
        # 分离题干和选项
        lines = question_content.split('\n')
        question_stem = ""
        options = []
        
        # 查找题干部分
        for i, line in enumerate(lines):
            if line.strip().startswith('A.'):
                # 从这里开始是选项
                question_stem = '\n'.join(lines[:i]).strip()
                options = lines[i:]
                break
        
        # 如果没有找到选项，整个作为题干
        if not options:
            question_stem = question_content
        
        choice_questions.append({
            'number': question_num,
            'type': '单选题',
            'stem': question_stem,
            'options': options
        })
    
    # 提取应用题
    app_pattern = r'(\d+)\.\s*(.*?)(?=\n\d+\.|\n##|$)'
    app_matches = re.findall(app_pattern, application_questions_text, re.DOTALL)
    
    application_questions = []
    for match in app_matches:
        question_num = match[0]
        question_content = match[1].strip()
        
        application_questions.append({
            'number': question_num,
            'type': '综合应用题',
            'content': question_content
        })
    
    return choice_questions, application_questions


def create_word_document(choice_questions, application_questions, year, output_path):
    """
    创建Word文档
    """
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f'{year}年数据结构考研真题及解析')
    title_run.font.name = '黑体'
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 添加选择题部分
    if choice_questions:
        choice_section = doc.add_heading(level=1)
        choice_section_run = choice_section.add_run('一、单项选择题')
        choice_section_run.font.name = '黑体'
        choice_section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        choice_section_run.bold = True
        
        for q in choice_questions:
            # 题号和题干
            para = doc.add_paragraph()
            run = para.add_run(f"{q['number']}. {q['stem']}")
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 选项
            for option in q['options']:
                option_para = doc.add_paragraph()
                option_run = option_para.add_run(option.strip())
                option_run.font.name = '宋体'
                option_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加空白行
            doc.add_paragraph()
    
    # 添加应用题部分
    if application_questions:
        app_section = doc.add_heading(level=1)
        app_section_run = app_section.add_run('二、综合应用题')
        app_section_run.font.name = '黑体'
        app_section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        app_section_run.bold = True
        
        for q in application_questions:
            para = doc.add_paragraph()
            run = para.add_run(f"{q['number']}. {q['content']}")
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加空白行
            doc.add_paragraph()
    
    # 添加解析部分（留空供手动填写）
    solution_section = doc.add_heading(level=1)
    solution_section_run = solution_section.add_run('题目解析')
    solution_section_run.font.name = '黑体'
    solution_section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    solution_section_run.bold = True
    
    # 为每道题添加解析预留空间
    all_questions = choice_questions + application_questions
    for q in all_questions:
        para = doc.add_paragraph()
        run = para.add_run(f"第{q['number']}题解析：")
        run.font.name = '宋体'
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加解析的空白区域
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
    
    doc.save(output_path)
    print(f"成功生成Word文档: {output_path}")


def process_all_years():
    """
    处理所有年份的文件
    """
    # 获取所有数据结构年份文件
    md_files = []
    for file in os.listdir('.'):
        if file.startswith('数据结构') and file.endswith('.md'):
            year = file[4:8]  # 提取年份
            if year.isdigit():
                md_files.append((file, year))
    
    # 确保输出目录存在
    os.makedirs('输出文档', exist_ok=True)
    
    for file, year in md_files:
        print(f"正在处理 {year} 年的数据...")
        
        try:
            choice_questions, application_questions = extract_questions_from_md(file)
            output_path = f"输出文档/{year}_数据结构题目及解析.docx"
            
            create_word_document(choice_questions, application_questions, year, output_path)
        except Exception as e:
            print(f"处理 {year} 年文件时出现错误: {str(e)}")


def process_single_year(year):
    """
    处理单个年份
    """
    file_name = f"数据结构{year}.md"
    
    if not os.path.exists(file_name):
        print(f"文件 {file_name} 不存在")
        return
    
    print(f"正在处理 {year} 年的数据...")
    
    try:
        choice_questions, application_questions = extract_questions_from_md(file_name)
        output_path = f"输出文档/{year}_数据结构题目及解析.docx"
        
        # 确保输出目录存在
        os.makedirs('输出文档', exist_ok=True)
        
        create_word_document(choice_questions, application_questions, year, output_path)
    except Exception as e:
        print(f"处理 {year} 年文件时出现错误: {str(e)}")


if __name__ == "__main__":
    # 安装依赖命令: pip install python-docx
    print("开始处理数据结构历年真题...")
    
    # 处理特定年份（例如2009年）
    # process_single_year(2009)
    
    # 或处理所有年份
    process_all_years()